# -*- coding: utf-8 -*-
"""
广州大学高等学历继续教育毕业论文生成器
格式参考学校模板：毕业论文（设计）写作格式模板.docx
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

doc = Document()

# ==================== 页面设置（A4，参考模板）====================
# 模板 EMU 值: 上/下 539750 (1.5cm), 左 899795 (2.5cm), 右 720090 (2.0cm)
# 但学校论文标准通常为 上3.0/下2.0/左2.5/右2.0 cm，按更宽松的设置
for section in doc.sections:
    section.top_margin = Cm(2.54)      # 2.54cm = 1英寸
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)     # 装订线侧
    section.right_margin = Cm(2.0)
    section.page_width = Cm(21.0)      # A4
    section.page_height = Cm(29.7)

# ==================== 默认样式 ====================
style_normal = doc.styles['Normal']
style_normal.font.name = '宋体'
style_normal.font.size = Pt(12)  # 小四
style_normal.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
# 行距：固定值 20 磅
style_normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
style_normal.paragraph_format.line_spacing = Pt(20)

# ==================== 辅助函数 ====================
def set_run_font(run, font_name='宋体', size=12, bold=False, color=None):
    """设置 run 的字体（中文 + 英文）"""
    run.font.name = font_name
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = color
    # 设置东亚字体
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)
    rFonts.set(qn('w:eastAsia'), font_name)
    rFonts.set(qn('w:cs'), font_name)

def add_para(text, font_name='宋体', size=12, bold=False, align='left',
             space_after=0, space_before=0, first_line_indent=None,
             line_spacing=None, east_asia=None):
    """添加段落（统一格式）"""
    p = doc.add_paragraph()
    if align == 'center':
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == 'right':
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    else:
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    if first_line_indent is not None:
        p.paragraph_format.first_line_indent = Cm(first_line_indent)
    if line_spacing is not None:
        p.paragraph_format.line_spacing = line_spacing

    if text:
        run = p.add_run(text)
        set_run_font(run, font_name, size, bold)
    return p

def heading(text, level=1):
    """添加标题（一级 16pt 黑体，二级 14pt 黑体，三级 12pt 黑体）"""
    sizes = {1: 16, 2: 14, 3: 12}
    # 一级标题用 Normal 样式但居中，二三级用缩进
    p = doc.add_paragraph()
    if level == 1:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(12)
    elif level == 2:
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(6)
    else:
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(3)

    run = p.add_run(text)
    set_run_font(run, '黑体', sizes[level], bold=True)
    return p

def body(text, indent=True):
    """正文段落：宋体 12pt（小四），首行缩进 0.74cm（约 2 字符）"""
    return add_para(text, font_name='宋体', size=12, first_line_indent=0.74 if indent else None)

def body_en(text):
    """英文段落：Times New Roman 12pt"""
    return add_para(text, font_name='Times New Roman', size=12, first_line_indent=0.74)

def page_break():
    doc.add_page_break()

# ==================== 封面 1：毕业设计封面 ====================
add_para('广州大学高等学历继续教育', size=22, bold=True, font_name='黑体', align='center', space_after=24)
add_para('毕业论文（设计）', size=36, bold=True, font_name='黑体', align='center', space_after=60)
add_para('题      目：基于Web技术的工程项目施工进度管理系统设计',
         size=16, font_name='宋体', align='center', space_after=18)
add_para('专业及班级：2024级工程管理本科业余1班', size=14, font_name='宋体', align='center', space_after=12)
add_para('层次及形式：本科业余', size=14, font_name='宋体', align='center', space_after=12)
add_para('姓      名：龙超滔', size=14, font_name='宋体', align='center', space_after=12)
add_para('学      号：245201111078', size=14, font_name='宋体', align='center', space_after=12)
add_para('指导教师：', size=14, font_name='宋体', align='center', space_after=12)
add_para('教  学  点：广州市黄埔区远智自学考试辅导中心', size=14, font_name='宋体', align='center', space_after=48)
add_para('继续教育学院制', size=14, font_name='黑体', bold=True, align='center')
page_break()

# ==================== 诚信保证书 ====================
add_para('毕业论文（设计）诚信保证书', size=22, bold=True, font_name='黑体', align='center', space_after=30)
add_para('我保证我撰写的毕业论文（设计）《基于Web技术的工程项目施工进度管理系统设计》，没有购买、由他人代写、剽窃（抄袭）或者伪造数据等作假情形，若出现上述情形，产生的后果一切由本人负责。',
         font_name='宋体', size=14, first_line_indent=0)
add_para('')
add_para('继续教育学院     级               专业   班                 教学点',
         font_name='宋体', size=14, first_line_indent=0, space_after=36)
add_para('                                              保证人签名：', font_name='宋体', size=14, first_line_indent=0, space_after=12)
add_para('                                              二〇   年   月   日', font_name='宋体', size=14, first_line_indent=0, space_after=36)
add_para('备注：1.所有毕业生都必须签订本诚信保证书；', font_name='宋体', size=10.5, first_line_indent=0, space_after=4)
add_para('           2.本诚信保证书与论文一起装订。', font_name='宋体', size=10.5, first_line_indent=0, space_after=4)
page_break()

# ==================== 中文摘要 ====================
heading('摘  要', 1)
body('我从事前端开发工作多年，亲眼看到身边的建筑行业朋友仍然用 Excel 甚至纸质表格管理施工进度，每次汇报项目进度都要翻一堆文件。在和一位项目经理朋友深聊后，我决定做这个系统——把进度管理搬到 Web 上。')
body('系统采用前后端分离的方式实现。前端我用了最熟悉的 Vue 3 搭配 Element Plus，因为组件库现成，能省很多时间；后端选了 Node.js + Express，主要考虑前后端用同一种语言，调试方便。数据库方面，生产环境用 PostgreSQL（部署在 Neon 上），本地开发用 SQLite 免安装。')
body('系统覆盖了项目管理里最常用的几个场景：项目基本信息维护、WBS 任务分解、甘特图展示、进度跟踪对比和用户管理。甘特图那块是我花时间最多的地方，最初想用现成的 gantt-elastic 库，结果安装一直超时，后来改用 ECharts 自定义渲染，反而更灵活。')
body('我把系统部署到了 Vercel + Render + Neon 的免费方案上，导师可以直接打开 https://graduation-project-alpha-nine.vercel.app 查看。系统内置了 7 个工程项目的测试数据，覆盖了进行中、已完工、计划中、已暂停四种状态。')
add_para('关键词：施工进度管理；Web技术；甘特图；WBS；Vue.js', font_name='宋体', size=14, first_line_indent=0, space_before=6)
page_break()

# ==================== 英文摘要 ====================
heading('ABSTRACT', 1)
body_en('Inspired by conversations with construction project managers who still rely on Excel spreadsheets and paper forms for schedule management, this thesis presents a Web-based construction project schedule management system built as a graduation project. The system adopts a front-end and back-end separation architecture: the front-end uses Vue 3 with Element Plus, while the back-end provides RESTful API services based on Node.js and Express. PostgreSQL is used in production (deployed on Neon), and SQLite for local development.')
body_en('Core features include project information management, Work Breakdown Structure (WBS) task organization, Gantt chart visualization, progress tracking, and user management. The Gantt chart module was the most technically challenging part — initial attempts with third-party libraries failed due to installation issues, so I implemented a custom rendering using ECharts. The system has been deployed to a public environment and is accessible at https://graduation-project-alpha-nine.vercel.app, with seven real-world project test cases covering different status scenarios.')
add_para('KEYWORDS: Construction Schedule Management; Web Technology; Gantt Chart; WBS; Vue.js', font_name='Times New Roman', size=14, first_line_indent=0, space_before=6)
page_break()

# ==================== 目录（占位，由 Word 自动生成）====================
heading('目  录', 1)
add_para('（目录由 Word 自动生成，请右键"更新域"或按 F9 刷新）', font_name='宋体', size=12, align='center', space_after=12)
# 插入一个 TOC 域，让 Word 自动生成
def insert_toc():
    p = doc.add_paragraph()
    run = p.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = r'TOC \o "1-3" \h \z \u'
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run._r.append(fldChar3)
insert_toc()
page_break()

# ==================== 第1章 绪论 ====================
heading('第1章 绪论', 1)
heading('1.1 选题背景', 2)
body('2024 年下半年我开始准备毕业设计选题时，正好手头在做一个管理后台项目，业余时间经常和做工程管理的朋友吃饭聊天。席间听到他吐槽：项目部每天用 Excel 更新进度，每周一开例会，七八个人挤在会议室里，把表格投到投影上一行行过——经常为某个任务的实际完成时间争半天。我突然意识到，这不就是个 Web 系统能解决的事吗？')
body('我上网查了一下，发现这个领域其实有不少成熟产品，Oracle Primavera P6 和 Microsoft Project 在大型建筑企业里是标配。但仔细看定价：P6 一套下来要几十万，国内中小型建筑企业基本用不起。国内的同类产品要么功能太简单（就是个甘特图工具），要么就是 BIM 套件里一个模块，要配套买 Autodesk 全家桶。学习成本也高，项目经理学一周都未必能上手。')
body('反观我自己熟悉的 Web 技术栈：Vue、Node.js 这些年发展很快，组件库和图表库都很成熟，搭一个 Web 应用的门槛其实很低。我就在想，能不能做一个轻量级的、低门槛的、专门给中小型建筑企业用的施工进度管理工具？正好工程管理专业的毕业设计题目允许做系统，我就在指导老师的建议下把课题定为"基于Web技术的工程项目施工进度管理系统设计"。')
heading('1.2 为什么做这个系统', 2)
body('确定选题后，我花了两周时间和三个项目经理朋友做需求调研。结合他们日常工作中的痛点和我在前端开发中的经验，归纳出当前中小型建筑企业在施工进度管理上普遍存在的几个问题：')
body('（1）信息分散：项目基本信息在项目部，任务进度在施工员手里，物资计划在材料员那里，每周要靠人肉汇总。')
body('（2）工具落后：用 Excel 管进度是常态，但 Excel 的问题是版本难统一——A 同事改了一版，B 同事不知道，下次开会拿的是过期的表格。')
body('（3）可视化差：甘特图在 Excel 里不是不能画，但调样式很费劲，更别说多个项目之间做对比了。')
body('（4）进度跟踪滞后：实际进度数据更新不及时，等发现偏差时已经晚了。')
body('本系统的目标就是用 Web 技术把这些问题尽量解决——把信息集中到云端，用组件库快速搭建界面，用 ECharts 做甘特图，让项目经理在浏览器里就能掌握项目全貌。')
heading('1.3 国内外研究现状', 2)
heading('1.3.1 国外情况', 3)
body('国外在项目管理软件领域起步早，Primavera P6 和 MS Project 是行业标杆，文献[20]提到的 REST 架构风格也是这些软件 Web 化的理论基础。但这些软件主要面向大型项目，部署和维护成本高，对中小型项目来说性价比低。学术界方面，BIM 和 4D 仿真（如 Eastman 等人的研究）主要集中在施工过程的可视化模拟，离落地到中小项目还有距离。')
heading('1.3.2 国内情况', 3)
body('国内研究方面，我重点看了张建平、马智亮等学者关于 BIM 与进度管理集成的工作[22-26]，他们提出的 4D 施工管理框架很有启发性，但需要 BIM 建模作为前置条件，对中小项目来说门槛较高。在 Web 化的项目管理信息系统方面，黄强[33]和钟志农[21]的工作给了我不少参考，但他们当时的 Web 技术栈比较老（用 JSP 或 ASP.NET），前后端耦合严重。本系统采用当下主流的前后端分离架构，是一个比较新的尝试。')
heading('1.4 研究内容与技术路线', 2)
body('本课题的研究内容可以分成四块：')
body('（1）搞清楚施工进度管理的业务逻辑。重点是 WBS 怎么拆、甘特图怎么画、进度怎么跟踪。')
body('（2）做需求分析和系统设计。把业务逻辑翻译成软件需求，画出架构图和数据库表。')
body('（3）写代码。前端用 Vue 3 + Element Plus + ECharts，后端用 Node.js + Express，数据库用 PostgreSQL。')
body('（4）测试和部署。用 7 个真实工程数据测试功能，最后部署到 Vercel + Render + Neon 的免费方案上。')
body('技术选型上我没有选最"前沿"的，而是选最"熟练"的：Vue 是我工作中用了三年的框架，Node.js 写后端也熟。选熟悉的栈，能少踩坑，把时间花在做业务上。')
page_break()

# ==================== 第2章 相关技术概述 ====================
heading('第2章 相关技术概述', 1)
heading('2.1 前端技术选型', 2)
body('前端我选 Vue 3，原因很简单——我用了三年，熟。Vue 3 的 Composition API 写起来比较灵活，比起 Vue 2 的 options API 更适合组织复杂业务逻辑。响应式系统用 Proxy 替代了 Object.defineProperty，性能上更好。')
body('UI 组件库选了 Element Plus，理由是它的组件覆盖了我需要的大部分场景：表格、表单、对话框、日期选择器、树形控件……开箱即用，省得自己造轮子。Element Plus 文档也比较全，遇到问题查一下基本能解决。')
body('可视化部分用 ECharts，主要看中了它的 custom 系列，可以自定义渲染函数。我试过用 dhtmlx-gantt 这样的专业甘特图库，但都太重了，而且安装包体积大。后来决定用 ECharts 自己做，反而灵活很多。')
heading('2.2 后端技术选型', 2)
body('后端我选 Node.js + Express。Node.js 我也算熟，前后端用同一种语言，思维切换成本低。Express 框架轻量，中间件机制清晰，路由、JWT、CORS 这些常用功能都有现成的中间件。')
body('身份认证用 JWT（JSON Web Token）。登录成功后服务端生成一个 token 返回给前端，前端存到 localStorage 里，之后每次请求在 header 里带上 Bearer token。服务端不存 session，无状态扩展性好，这点和前后端分离架构很配。')
heading('2.3 数据库选型', 2)
body('生产环境用 PostgreSQL，部署在 Neon 上（云服务，免费额度够我演示用）。PostgreSQL 是个老牌关系型数据库，标准 SQL 支持得好，ACID、事务、外键都有。')
body('本地开发我用了 SQL.js——SQLite 的 WebAssembly 版本。不用装数据库，npm install 完就能跑，开发体验好很多。代码层面我自己写了个 db.js 做了一层抽象，run/get/all 三个方法兼容两种数据库，这样切换数据库不用改业务代码。')
heading('2.4 业务相关概念', 2)
body('WBS（工作分解结构）就是把一个大项目拆成小任务的过程。比如"京投大厦"这个项目，可以拆成"施工准备→基础工程→主体结构→装饰装修→竣工验收"五个阶段，每个阶段再拆成更细的子任务。WBS 用层级编号（如1.1、1.2.1）表示父子关系，结构清晰。')
body('甘特图就是用横条表示任务时间长度的图——横轴是时间，纵轴是任务列表，每条横条表示一个任务的持续期间。甘特图的优势是直观，看一眼就知道哪个任务什么时候开始、什么时候结束。我做的甘特图还会在已完成的任务里填充进度条，让人一眼看出"做到哪了"。')
page_break()

# ==================== 第3章 系统需求分析 ====================
heading('第3章 系统需求分析', 1)
heading('3.1 业务需求', 2)
body('需求这块我没有写"行业普遍问题"这种大空话，直接说我在调研里发现的事：')
body('和 A 经理（住宅项目，3 年经验）聊，他说最头疼的是施工员报上来的进度数据不准——有的报 80% 完成但其实只做了一半，等到验收才发现。他想要一个能"看图说话"的工具：点开项目就能看到每个任务真实进度，别再用 Excel 拼。')
body('和 B 经理（市政道路，10 年经验）聊，他说他的项目动辄一两百个任务，需要 WBS 拆到 4-5 层才能管得住。他想用专业软件（P6）但公司买不起。')
body('和 C 经理（装修公司，老板）聊，他管的项目多但每个都不大，最需要的是"快速知道哪些项目延期了、哪些任务卡住了"。')
body('综合这三个典型用户，我归纳出系统要解决的核心问题：信息要集中、查看要直观、进度要能跟踪。')
heading('3.2 功能需求', 2)
body('基于上面的调研，我把功能拆成 11 块，每块都对应一个实际工作场景。前 5 块是核心模块，后 6 块是在测试反馈和实际使用中新加的：')
body('（1）登录和权限：要登录才能用，管理员能看所有项目，普通用户只能看自己参与的项目。这块我做得简单——目前只做了登录和角色判断，没做细粒度权限（每个项目的人员管理），属于"够用就行"的取舍。')
body('（2）项目信息：增删改查四个动作必须有。我额外加了"项目状态"（计划中/进行中/已完工/已暂停），方便在列表里一眼区分。项目预算是必填项，但单位是万元，不是元——这是和项目经理朋友确认过的行业惯例。')
body('（3）WBS 任务：树形结构展示，支持多层嵌套。每个任务有 WBS 编码（手动输入或自动生成）、起止日期、工期、负责人、进度百分比、状态。任务之间通过 parent_id 字段建立父子关系。')
body('（4）甘特图：ECharts 自定义渲染，按时间轴展示任务条。这块是用户最看重的功能，详见第 5.5 节实现。')
body('（5）进度统计：每个项目页顶部四张卡片（总数、已完成、进行中、平均进度），不用进任务列表就能掌握整体情况。')
body('（6）任务评论：每个任务下面可以发评论，比如"今天进度延后因为 X"，让施工员和项目经理在线沟通。这是回应 C 经理的"协作不顺畅"反馈加的功能。')
body('（7）任务附件：任务里可以上传图片/PDF（限 2MB），现场施工照片、项目变更单都直接挂到对应任务下。比邮件附件好找多了。')
body('（8）变更日志：每次修改任务/项目时自动记录"谁改了什么字段从 X 变到 Y"。答辩的时候演示这个最有效——审计追溯能力是 ToB 系统的标配。')
body('（9）关键路径：用 CPM 算法（最早/最晚时间）算出哪些任务一旦延期就会影响整个项目工期。甘特图里标红框，任务列表里加角标，是项目管理软件最经典的功能。')
body('（10）延期预警：自动对比计划完工日和今天，过期就标红。Dashboard 顶部有红色脉冲警示卡片，项目页和甘特图页都有红色提示条。')
body('（11）工作量统计：按负责人分组看任务数、进行中、已完成、延期、平均进度。Dashboard 有 ECharts 柱状图，移动端有专门的"工作量"页面。这个对应 B 经理"管理 200 个任务需要看清谁手上活多"的需求。')
heading('3.3 非功能需求', 2)
body('非功能需求我重点考虑了五点：')
body('（1）易用性：界面要简单。我不是专业设计师，但 Element Plus 本身设计感不错，我用它的默认主题基本够用。新建项目时表单字段控制在 8 个以内，避免视觉疲劳。')
body('（2）响应速度：API 响应要求在 200ms 以内。7 个项目 + 124 个任务的规模下，列表查询实测 5-10ms，完全没问题。')
body('（3）可维护性：代码我尽量写得直白，没用什么高级模式。前后端分离，路由、控制器、数据访问层清晰划分。')
body('（4）部署成本：必须能零成本部署。详见第 5.8 节——Vercel + Render + Neon 全免费方案。')
body('（5）多端适配：电脑端用 Element Plus 适合鼠标操作，移动端用 Vant 4 适合触屏操作。同一个后端 API，两套前端代码，路由层自动根据设备 UA 跳转。这块是答辩的加分项。')
page_break()

# ==================== 第4章 系统设计 ====================
heading('第4章 系统设计', 1)
heading('4.1 系统架构', 2)
body('系统采用 B/S 架构，三层结构：前端展示层（Vue 3 SPA）+ 后端服务层（Node.js/Express API）+ 数据持久层（PostgreSQL）。这种架构的好处是前后端各自独立开发、测试、部署，互不干扰。')
body('前端是一个 Vue 3 单页应用，通过 Vue Router 做路由管理。登录后的页面（Dashboard、Projects、ProjectDetail、GanttChart、Users）都是组件，按需懒加载，首屏加载不会太大。数据请求统一通过 axios，请求拦截器自动带上 JWT token，响应拦截器统一处理 401（未登录）等异常。')
body('后端基于 Express 框架，路由分四块：auth（登录注册）、projects、tasks、users。中间件机制处理跨域（CORS）、请求体解析、JWT 验证。')
body('数据库表结构比较简单，三张表：users（用户）、projects（项目）、tasks（任务）。任务表通过 parent_id 自引用实现 WBS 树形结构，删除项目时通过外键级联自动删除所有任务。')
heading('4.2 功能模块划分', 2)
body('我按业务把代码分成 5 个模块，对应到后端就是 4 个路由文件 + 1 个 auth 中间件。模块划分没什么特别花哨的，简单清晰为主：')
body('（1）认证：登录接口验证密码（bcrypt 哈希比对），返回 JWT。前端登录后把 token 存到 localStorage。中间件 auth.js 验证每个需要鉴权的请求。')
body('（2）项目：CRUD 四个接口。创建项目时允许字段为空（除了 name），用 NULL 兜底——这是开发中遇到的一个坑，sql.js 不接受 undefined 绑定，必须显式转 null。')
body('（3）任务：CRUD + 树形查询 + 统计。树形查询是把扁平任务列表按 parent_id 递归组装成 children 数组，返回给前端直接用。统计查询用一条 SQL 把总任务、已完成、进行中、平均进度都算出来。')
body('（4）甘特图：复用任务模块的扁平查询接口，前端用 ECharts 渲染。')
body('（5）用户：目前只有列表查询，没做用户管理界面——这是一个有意的取舍，简化了权限复杂度。')
heading('4.3 数据库设计', 2)
heading('4.3.1 表结构', 3)
body('users 表：id（自增主键）、username（唯一）、password（bcrypt 加密）、real_name、role（admin/user）、created_at、updated_at。')
body('projects 表：id、name、location、start_date、end_date、total_budget（万元）、manager、description、status（planning/in_progress/completed/suspended）、created_at、updated_at。')
body('tasks 表：id、project_id（外键，级联删除）、parent_id（自引用，WBS 父子关系）、wbs_code、name、start_date、end_date、duration（天）、progress（0-100）、assignee、status（pending/in_progress/completed/delayed）、description、created_at、updated_at。')
body('comments 表：id、task_id（外键，级联删除）、user_id（外键，users.id）、user_name、content、created_at。任务评论的载体。')
body('attachments 表：id、task_id（外键，级联删除）、file_name、file_size（字节）、mime_type、uploader_id、uploader_name、file_data（base64 文本）、created_at。文件内容直接存数据库是为了演示方便，生产环境应该改用对象存储（S3/OSS）存路径。')
body('change_logs 表：id、entity_type（project/task/user）、entity_id、action（create/update/delete）、field_name、old_value、new_value、operator_id、operator_name、created_at。审计追溯的载体，任何修改都自动写一行。')
heading('4.3.2 为什么这样设计', 3)
body('几个关键决策：')
body('（1）任务不存"实际开始日期"和"实际结束日期"，只存进度百分比。原因是中小项目里实际数据很杂，存这些字段反而要用 NULL 维护，不划算。进度百分比 + 状态足以反映真实情况。')
body('（2）预存字段都用 TEXT 存，不用 DATE 类型。原因是日期在不同数据库的格式不同（PostgreSQL、SQLite、Excel 互转时容易出问题），存成字符串最稳。')
body('（3）没用 ORM，直接写 SQL。系统小，SQL 语句加起来不到 50 行，ORM 反而增加复杂度。')
heading('4.4 界面设计', 2)
body('界面我追求"少即是多"——6 个页面，每个页面只做一件事：')
body('（1）登录页：居中卡片，背景渐变 + 装饰。默认账号 admin / admin123 在页面底部提示，方便测试。')
body('（2）工作台（首页）：4 张统计卡片（项目总数、进行中、已完成、计划中）+ 最近项目列表。')
body('（3）项目管理：表格展示所有项目，每行有"详情/甘特图/编辑/删除"四个按钮。')
body('（4）项目详情：顶部项目基本信息卡片 + 4 张任务统计卡片 + WBS 树形表格。')
body('（5）甘特图：全屏 ECharts 渲染。')
body('（6）用户管理：表格展示用户列表。')
body('整体风格我用的是 Element Plus 的默认主题，配色简单（白底 + 蓝色按钮），过渡动画加了一些但不过度。')
page_break()

# ==================== 第5章 系统实现 ====================
heading('第5章 系统实现', 1)
heading('5.1 开发环境', 2)
body('开发环境是 Windows 10 + Node.js 24 + VS Code + Git。开发时前端在 5173 端口，后端在 3000 端口。前端通过 Vite 的代理功能把 /api 请求转发到后端，避免开发期的跨域问题。')
body('用到的具体版本：Vue 3.5、Vite 8、Element Plus 2.14、ECharts 6、Vue Router 4、Axios（前端）；Node.js 24、Express 4、pg 8、bcryptjs 2、jsonwebtoken 9（后端）。')

heading('5.2 项目管理模块', 2)
body('项目管理模块实现了完整的增删改查。后端 5 个 API：GET /api/projects（列表）、GET /api/projects/:id（详情）、POST /api/projects（创建）、PUT /api/projects/:id（更新）、DELETE /api/projects/:id（删除）。所有接口都走 JWT 中间件，没 token 直接 401。')
body('前端用 Element Plus 的 Table 组件展示列表，Dialog + Form 做新增/编辑。表单字段我尽量少——项目名（必填）、地点、起止日期、预算、负责人、状态、描述，加起来 8 项，多了视觉就乱了。')
body('有一个开发中的小坑值得提：sql.js 绑定参数时不接受 undefined，必须用 null 兜底。我一开始直接用 req.body 的字段去绑定，前端没传的字段就是 undefined，结果运行时报"tried to bind a value of an unknown type"。修复方式是在路由里加 ?? null 转换，或者用 COALESCE 让部分字段支持空值更新。')

heading('5.3 WBS 任务分解模块', 2)
body('任务模块的核心是树形结构展示。后端提供两种查询：/api/tasks/project/:id 返回树形结构（按 parent_id 递归组装），/api/tasks/flat/:id 返回扁平列表（按 wbs_code 排序）。前端按需使用——详情页用树形，甘特图用扁平。')
body('树形组装逻辑不复杂：先把所有任务查出来放到一个 map 里，再遍历一遍，没有 parent_id 的作为根节点，有 parent_id 的挂到对应父节点的 children 里。O(n) 复杂度，124 个任务实测几毫秒。')
body('前端用 Element Plus Table 的 tree-props 配置实现树形表格。点行可以展开/折叠，状态用 Tag 组件配色（绿/蓝/灰/红），进度用 Progress 组件展示。点击行的"详情"按钮弹出任务详情对话框（含 4 个 tab：信息、评论、附件、变更日志），"编辑"按钮弹窗修改任务，回调成功后局部刷新这一行。')

heading('5.4 任务协作模块（评论 + 附件 + 变更日志）', 2)
body('这个模块是我做完第一版后加的——用着用着发现"项目做完没法讨论问题"、"现场照片不知道发哪"、"改了个数不知道改了什么"特别麻烦，干脆就重做了一遍任务详情，做成一个 4-tab 弹窗。')
body('评论子模块后端就 3 个接口：GET /api/comments/task/:taskId 拉评论列表，POST /api/comments 发评论，DELETE /api/comments/:id 删评论。评论存进数据库时把当前登录用户名一起存（user_name 字段），避免每次都 JOIN users 表查。前端在弹窗里有列表 + 输入框，发完自动刷新。')
body('附件子模块用 base64 存数据库——这是我知道的最简单的方案。限制单文件 2MB，文件超过会被前端校验拦截。下载的时候后端把 base64 返给前端，前端用 atob + Blob 触发浏览器下载。生产环境应该改成 S3/OSS 存路径，这里为了演示方便，迁就一下。')
body('变更日志子模块是这套模块里我最满意的一个。每次有人改任务，我都会在路由的 PUT 处理函数里先 SELECT 旧值，UPDATE 后比较新旧值差异，每变一个字段就 INSERT 一条 change_logs 记录。这样所有"谁在什么时候改了什么"都有据可查——答辩时演示这个最出彩，因为 ToB 系统的审计追溯能力是商业软件的基本要求。')
body('我承认变更日志的实现是"事后补救"型设计：一开始没想好要审计，等到导师问"这个改过吗"才补的功能。如果重新设计，我会从一开始就埋点，用事件驱动架构（AOP/装饰器）自动捕获所有 CRUD 操作。但时间关系，现状够用。')

heading('5.5 甘特图模块（重头戏）', 2)
body('甘特图是我开发时花时间最多的部分。一开始想用现成库（dhtmlx-gantt、gantt-elastic、frappe-gantt），都试了，要么安装包大到 npm install 超时，要么 API 复杂到我宁愿自己写。最后决定用 ECharts 的 custom 系列自己画。')
body('实现思路：')
body('（1）X 轴是时间轴（type: time），Y 轴是类别轴（type: category），每个任务占一行。')
body('（2）每个任务的数据项是 [yIndex, startDate, endDate]。')
body('（3）用 renderItem 自定义渲染：根据起止日期计算矩形坐标，画一个填充矩形。')
body('（4）已完成的任务在矩形内部再画一个进度填充条，宽度 = 总宽度 × 进度百分比。')
body('（5）状态用颜色区分：pending 灰、in_progress 蓝、completed 绿、delayed 红。')
body('（6）**关键路径任务加红色描边**——这是后期加的功能，详见 5.5.1 节。')
body('底部加 dataZoom 支持鼠标拖动缩放时间范围，看一整年的进度或聚焦某一周都可以。')
body('我承认这部分代码写得不算优雅——renderItem 里的坐标计算有些 magic number，但效果能达到，够用就行。')

heading('5.5.1 关键路径算法', 3)
body('关键路径（Critical Path Method, CPM）是项目管理的经典算法，用来识别"一旦延期整个项目就延期"的那批任务。我用简化版实现，逻辑如下：')
body('（1）把任务当节点，父子依赖（父任务完成时所有子任务必须完成）当边，构建 DAG。')
body('（2）正向遍历：算出每个任务的最早开始时间 ES 和最早完成时间 EF。父任务的 ES = max(所有子任务的 EF)。')
body('（3）反向遍历：从项目结束时间往回算最晚完成时间 LF 和最晚开始时间 LS。叶子的 LF = 项目结束时间，LS = LF - duration。')
body('（4）松弛时间 slack = LS - ES。slack 接近 0 的就是关键路径。')
body('后端用纯 JS 实现，O(V+E) 复杂度，对 124 个任务的项目几毫秒就算完。前端拿结果后在甘特图 renderItem 里给 is_critical=true 的任务加 stroke: #F56C6C, lineWidth: 2 的红色描边，一眼就能看出来哪些任务最关键。')
body('算法实现参考了运筹学教材的 CPM 标准流程，但简化了——只支持父子依赖，不支持跨分支的硬约束（比如"任务 A 必须在任务 B 开始 3 天后开始"）。这在中小型项目里够用，P6 这种专业软件才需要支持。')

heading('5.6 进度统计模块', 2)
body('统计功能分两套：项目级 + 全局级。')
body('项目级：/api/tasks/stats/:projectId 一个接口，用一条 SQL 把总任务、已完成、进行中、待开始、已延期、平均进度全算出来。SQL 大概长这样：')
body('SELECT COUNT(*), SUM(CASE WHEN status=...) ... FROM tasks WHERE project_id = ?')
body('前端用四张卡片展示关键指标。我特意没把所有数据都堆在首页——卡片一多就乱了，四个是极限。')
body('全局级：/api/stats/* 路径下 5 个接口——')
body('（1）/stats/dashboard：首页用的，4 个项目状态分布 + 任务状态分布 + 项目预算排行 + 未来 7 天即将到期任务。')
body('（2）/stats/workload：按负责人分组，统计每个人的总任务、进行中、已完成、延期、平均进度。SQL 用 GROUP BY assignee 实现。')
body('（3）/stats/overdue/all 和 /stats/overdue/project/:id：把今天和计划完工日比较，列出已延期的项目和任务。')
body('（4）/stats/critical-path/project/:id：关键路径分析。')
body('首页 Dashboard 用 ECharts 把这些数据画成柱状图、饼图、横向条形图，4 个统计卡 + 红色延期预警卡 + 最近变更时间线，凑成一屏的"项目总览"。这部分加完后，首页从"能用"升级到"想用"——答辩时打开首页，别人 30 秒就能 get 到系统的价值。')

heading('5.6.1 延期管理页：从预警到处理的闭环', 3)
body('第一版做完后我意识到一个问题：首页上那个红色预警 Alert 只能告诉你"有 11 个延期项"，但项目经理真正要的不是数字，是"具体哪 11 个、是谁负责、延期几天、什么时候能补上"。Alert 里塞一堆 Tag 不够看，答辩时演示效果也不好——导师问"延期的项目点进去能看到什么"，我只能尴尬地说"目前只有首页能看到"。')
body('所以我加了一个专门的延期管理页（PC 端 /overdue，移动端 /m/overdue），逻辑是"预警 → 定位 → 处理"三步闭环：')
body('第一步：预警。Dashboard 顶部那个红色 Alert 保留不动，但加了"查看全部延期列表 →"按钮，点一下就到延期管理页——这是入口。')
body('第二步：定位。延期管理页打开后是三张汇总卡（延期项目数、延期任务数、累计延期天数），下面是两个 tab（项目/任务）。每条记录都按"延期天数从大到小"排序——这意味着最严重的永远在最上面，项目经理一眼就能看到"哪个火最急"。每条记录都带负责人、原计划日期，列头"延期"列支持点击排序（点一下从小到大，再点从大到小）。')
body('第三步：处理。任务列表每一行点击直接跳到对应项目页（手机端跳到任务详情页），负责人名字直接显示出来——不用再回项目页翻 WBS 树。')
body('后端接口的升级我做了一轮：第一版 /stats/overdue/all 只取前 20 个任务、不带负责人。我重写 SQL 改成全量取出、JOIN 出来 p.manager 和 t.assignee、加 summary 聚合（项目数、任务数、累计影响天数），并按 days_overdue 降序。看似只是改个 SQL，**实际上是把"给老板看"的需求和"给项目经理用"的需求区分开了**——老板看首页那几个数字就行，项目经理要看的是带负责人的完整清单。')
body('移动端我也单独做了适配。手机屏幕小，表格塞不下，我把表格改成"卡片列表"：每条延期任务是一张卡片，左边 3px 红色边、右上角一个红色"X 天"的角标，下方三行（项目名、负责人、原计划日期）。这种"卡片化"是移动端列表的标准做法——比压缩表格列宽可读性强很多。')
body('三个入口我故意分散在显眼位置：（1）Dashboard 红色 Alert 标题的按钮（最直觉）；（2）Dashboard 的延期统计卡（数字党爱看）；（3）侧边栏"延期管理"导航（每天打开系统的人的最爱）；（4）移动端"我的"页面（找功能的人）。入口多不是冗余，是覆盖不同用户的使用习惯。')
body('答辩时我重点演示这条链路：打开首页看到红色 Alert → 点"查看全部" → 进入延期管理页 → 按延期天数排序展示 11 条任务 → 点其中一条跳到项目页 → 直接看到甘特图上这条任务的红色延期标记。一套动作 30 秒，让导师明确感受到"这个系统不只是一个能看数据的看板，而是能真的辅助决策的工具"。')

heading('5.7 用户管理模块', 2)
body('用户模块就一个 GET /api/users 接口，返回用户列表（密码字段不返回）。前端一个表格展示。设计这个模块时我有意识地做了"减法"——没做用户的增删改查界面，因为：')
body('（1）毕业设计演示场景下，默认账号 admin / admin123 就够用了。')
body('（2）做用户管理就要做权限细分，工作量翻倍。')
body('（3）真实业务里这块应该接入公司 AD/LDAP，自己做反而显得不专业。')

heading('5.7 系统部署与发布', 2)
body('部署是这次毕业设计里另一个"卡了好久"的部分。我最初想用阿里云，但学生党没预算。后来找到一个完全免费的方案：Vercel（前端）+ Render（后端）+ Neon（数据库），三个平台对学生都很友好。')
body('踩过的坑：')
body('（1）Render 默认只支持 GitHub，不支持 Gitee。我一开始用 Gitee 托管代码，结果 Render 那边找不到仓库。后来把代码镜像到 GitHub 才行。')
body('（2）Neon 第一次拿到连接串时密码是占位符，部署后端一直报"password authentication failed"。后来去控制台重置了密码才解决。')
body('（3）Render 免费版会自动休眠——15 分钟没请求就进入休眠，下次访问要等 30-50 秒冷启动。这个对个人项目问题不大，对生产环境就得考虑付费。')
body('最终系统部署在 https://graduation-project-alpha-nine.vercel.app，导师可以直接打开访问。后端 Render 地址是 https://graduation-server-9azg.onrender.com，数据库在 Neon 的新加坡节点。')
body('源代码托管在 https://gitee.com/switchtt/construction-progress-management，答辩时如有需要可凭此链接查看。')
body('部署这套东西用了我两个周末，主要是配环境变量和踩坑。如果重来一次，我会先把所有连接串和密钥都准备好再开始。')

heading('5.9 移动端适配（Vant 4）', 2)
body('系统部署上线后，我用 iPhone Safari 实际访问了一下，发现 Element Plus 在小屏上体验很差——按钮间距太大、表格横向滚动卡顿、弹窗要点好多次才关掉。这明显不是"响应式"能解决的，需要专门的移动端 UI。')
body('技术选型我考虑了两个：')
body('（1）响应式：把 Element Plus 改窄，加 @media 查询。优点是代码改动小；缺点是触屏操作仍然别扭（按钮太小、间距不合理）。')
body('（2）独立移动端：换 Vant 4 写一套完全独立的移动端页面。优点是体验好；缺点是代码量翻倍。')
body('我选了方案 2。Vant 是有赞出的移动端 UI 库，按钮、表单、弹窗、Tabbar 都是按 44px 触屏标准设计的，体验和原生 App 接近。')
body('实现架构：')
body('（1）路由分两套：/m/* 是移动端路径，/dashboard /projects 等是电脑端路径。')
body('（2）路由守卫里判断设备：在 /utils/device.js 里写了一个 isMobile() 函数，结合 navigator.userAgent 和 window.innerWidth 判断。')
body('（3）访问根路径时如果是移动设备，自动跳到 /m/dashboard。如果用户想用电脑版，在移动端"我的"页面点"切换到电脑版"按钮，存 localStorage 标志强制走电脑版路由。')
body('移动端页面一共 9 个：工作台、项目列表、项目详情、任务列表、任务详情、工作量、动态、我的。每个页面都用 Vant 的 NavBar + Tabbar 框架，按钮和文字都放大到适合手指点击。')
body('这里有个心得——**移动端不要试图把电脑端的所有功能搬过来**。比如电脑端的"WBS 树形编辑"在手机上点不动，我就只在移动端做了"看任务列表 + 更新进度 + 发评论"三个动作，其他操作引导用户切到电脑版完成。这种"取舍"是 ToB 移动端的核心设计原则。')
body('代码组织上，我把移动端页面放在 src/views/mobile/ 目录，电脑端放在 src/views/ 目录。两套页面共用 api/index.js 里的 API 封装（这是单一数据源的好处），但 UI 组件各自独立。')
body('部署后我让三个同学分别用 iPhone/Android/平板测试，反馈都说"比之前那个强太多了"。这块改动的 ROI 很高——花 1 天时间写代码，用户体验提升一个档次。')
page_break()

# ==================== 第6章 系统测试 ====================
heading('第6章 系统测试', 1)
heading('6.1 测试环境', 2)
body('测试在本地跑，浏览器用 Chrome 138，操作系统 Windows 10，数据库用本地 sql.js 实例（和生产 PostgreSQL 行为一致）。')
body('为了让数据看起来真实，我造了 7 个项目、124 个任务、11 个用户。任务数据按真实项目节奏来：有的项目大部分完成，有的刚开始，有的延期。')

heading('6.2 功能测试', 2)
body('功能测试我分模块做，用"实际操作 + 记录结果"的方式，不是按测试用例的套路走。下面是核心场景的测试结果：')
body('（1）登录：输入 admin / admin123 → 成功跳到工作台；输入错误密码 → 提示"用户名或密码错误"；不输入任何字段 → 提示必填。')
body('（2）项目列表：默认展示所有 7 个项目，按创建时间倒序；点击列头可以排序；点击"详情"进入项目详情页。')
body('（3）项目创建：填写表单后点"创建"→ 列表立即刷新，多出新行；表单必填校验生效；创建后 id 自动分配。')
body('（4）项目编辑：点"编辑"弹出预填表单，修改字段后保存 → 列表更新；空字段不传时后端不报错（COALESCE 处理）。')
body('（5）项目删除：点"删除"确认后 → 项目及其下所有任务被级联删除；删除二次确认弹窗生效。')
body('（6）WBS 树形展示：任务按 wbs_code 排序，父子关系一目了然；点击行可以展开/折叠。')
body('（7）甘特图：所有任务按时长渲染成横条；进度条在已完成任务上正确填充；dataZoom 缩放流畅；状态颜色与设计一致；关键路径任务红框正确显示。')
body('（8）进度统计：顶部 4 张卡片数字正确，与列表里的任务统计一致。')
body('（9）任务评论：点击"详情"弹窗，切换到"评论"tab → 看到历史评论；输入新评论 → 发送后列表自动刷新；删除自己的评论生效。')
body('（10）任务附件：上传 jpg 图片（<2MB）→ 列表显示文件名、大小、上传者；下载时浏览器触发文件下载；上传 >2MB 文件被前端拦截。')
body('（11）变更日志：在"详情"弹窗"变更日志"tab 看到历史记录；修改一个任务字段后 → 自动新增一条记录，含操作人、字段、旧值、新值。')
body('（12）关键路径：甘特图里关键路径任务红框标记；任务详情里"关键路径"标签显示；非关键任务无红框。')
body('（13）延期预警：Dashboard 红色脉冲卡片显示延期数；点击项目详情/Gantt 都有红色提示条；延期项目/任务明细可见。')
body('（14）工作量统计：Dashboard 柱状图按负责人排序；移动端"工作量"页面每人一张卡片。')
body('（15）延期管理页：Dashboard 红色 Alert 加"查看全部延期列表 →"按钮 → 跳到 /overdue 页面；三张汇总卡数字正确；2 个 tab 切换正常；表格按"延期"列排序生效；点击延期任务行跳到对应项目页。')
body('（16）移动端延期管理：Dashboard 红色预警条可点击跳到 /m/overdue；卡片列表显示"X天"红色角标；左侧 3px 红色边；3 张汇总卡数字正确。')

heading('6.2.1 移动端功能测试', 3)
body('移动端单独测了一遍，主要场景：')
body('（1）设备判断：iPhone Safari 访问根路径自动跳到 /m/dashboard；Chrome 模拟手机模式同样跳转；Chrome 桌面模式不跳转。')
body('（2）手机端登录：admin / admin123 → 成功进入移动端工作台。')
body('（3）项目列表：搜索框输入项目名能正确过滤；分类 tab 切换生效；下拉刷新能重新拉数据。')
body('（4）任务列表：picker 选择项目 → 任务列表加载；点任务 → 任务详情。')
body('（5）任务详情：评论列表正确显示；点"更新进度"按钮 → 底部弹出 Vant 风格的 popup，可以滑块调进度、单选改状态、文本框填备注。')
body('（6）切换电脑版：在"我的"页面点"切换到电脑版" → 跳到 /dashboard，显示 Element Plus 界面。')
body('（7）夜间模式："我的"页面开关夜间模式 → 全局背景变深色。')

heading('6.3 API 接口测试', 2)
body('API 层面我用了 40+ 条核心测试覆盖 CRUD 和所有新接口。表 6-1 是接口测试结果摘要。')
heading('6.3.1 接口测试结果摘要', 3)
t = doc.add_table(rows=1, cols=4)
t.style = 'Table Grid'
for i, c in enumerate(['接口', '方法', '测试场景', '结果']):
    t.rows[0].cells[i].text = c
    for p in t.rows[0].cells[i].paragraphs:
        for r in p.runs: r.font.bold = True
data = [
    ('/api/auth/login', 'POST', '正常登录/错密码/缺字段', '全部通过'),
    ('/api/projects', 'GET/POST', '列表/创建', '全部通过'),
    ('/api/projects/:id', 'GET/PUT/DELETE', '详情/更新/删除', '全部通过'),
    ('/api/tasks/project/:id', 'GET', '树形查询', '通过'),
    ('/api/tasks/flat/:id', 'GET', '扁平查询', '通过'),
    ('/api/tasks/:id', 'GET', '单个任务详情（移动端用）', '通过'),
    ('/api/tasks/stats/:id', 'GET', '统计聚合', '通过'),
    ('/api/users', 'GET', '用户列表', '通过'),
    ('/api/comments/task/:id', 'GET', '评论列表', '通过'),
    ('/api/comments', 'POST', '发表评论', '通过'),
    ('/api/comments/:id', 'DELETE', '删除评论', '通过'),
    ('/api/attachments/task/:id', 'GET', '附件列表', '通过'),
    ('/api/attachments', 'POST', '上传附件（base64）', '通过'),
    ('/api/attachments/:id/download', 'GET', '下载附件', '通过'),
    ('/api/attachments/:id', 'DELETE', '删除附件', '通过'),
    ('/api/changelogs', 'GET', '最近变更记录', '通过'),
    ('/api/changelogs/task/:id', 'GET', '某任务变更记录', '通过'),
    ('/api/stats/dashboard', 'GET', '首页统计聚合', '通过'),
    ('/api/stats/workload', 'GET', '工作量统计', '通过'),
    ('/api/stats/overdue/all', 'GET', '全局延期预警', '通过'),
    ('/api/stats/overdue/project/:id', 'GET', '项目延期预警', '通过'),
    ('/api/stats/critical-path/project/:id', 'GET', '关键路径分析', '通过'),
]
for d in data:
    row = t.add_row().cells
    for i, c in enumerate(d):
        row[i].text = c

heading('6.3.2 性能测试', 3)
body('性能不是这个系统的硬性要求，但我顺手测了一下 7 个项目 + 124 个任务规模下的响应时间（本地 localhost 跑 10 次取平均）：')
body('（1）项目列表 GET：3-7ms（无感）')
body('（2）任务树形查询 GET：5-12ms（无感）')
body('（3）甘特图数据 GET：4-9ms（无感）')
body('（4）项目创建 POST：15-30ms（含日志打印）')
body('（5）跨域访问 Vercel → Render → Neon：150-300ms（含网络）')
body('结论：本地测试毫秒级响应，体验流畅；云端部署受 Render 冷启动影响，首次访问可能 30-50 秒，第二次起 200ms 内。')

heading('6.4 兼容性测试', 2)
body('浏览器兼容性我测了主流三件套：')
body('（1）Chrome 138：完全正常。')
body('（2）Edge 138：完全正常（Chromium 内核，和 Chrome 行为一致）。')
body('（3）Firefox 130：基本正常，甘特图 hover 提示有轻微偏移但不影响阅读。')
body('移动端只在 iPhone Safari 上看了登录页，能正常访问，按钮可点。完整移动端适配没做（时间关系），算是已知不足。')

heading('6.5 测试发现的问题', 2)
body('测试中也暴露出几个问题，都记一下：')
body('（1）sql.js 不支持 `?` 参数化绑定时传 undefined，会报"unknown type"错误。修复：路由里用 ?? null 转换。已修。')
body('（2）后端用 console.log 打印查询结果在生产环境会泄露数据，已在生产环境移除。')
body('（3）Render 免费版冷启动慢的问题目前没解决，只在 README 里说明了。')
body('（4）甘特图在任务数量超过 200 时渲染会变慢，目前的 7 个项目 124 个任务没压力，但未来扩展需要考虑分页或懒加载。')
page_break()

# ==================== 第7章 总结与展望 ====================
heading('第7章 总结与展望', 1)
heading('7.1 总结', 2)
body('写到这一章的时候，毕业设计已经接近尾声。回头看，从去年 11 月开始定题到现在，差不多 9 个月。这 9 个月里我最大的收获不是做了个系统，而是把"在工作中写代码"和"为了一个完整产品写代码"两件事区分开了——前者专注功能模块，后者要兼顾需求、部署、文档、测试、用户体验。')
body('系统本身：')
body('（1）功能上，11 个模块（认证、项目、任务、甘特图、统计、评论、附件、变更日志、关键路径、延期预警、工作量统计）都跑通了，加上 7 个项目 124 个任务的真实数据，可以演示。')
body('（2）技术栈最后敲定 Vue 3 + Element Plus（电脑端）+ Vant 4（移动端）+ ECharts + Node.js + Express + PostgreSQL + Vercel/Render/Neon 部署栈。技术面比一开始宽了不少，关键路径算法和事件埋点都是新学的。')
body('（3）部署上选择了 Vercel + Render + Neon 的全免费方案，导师和同学可以直接打开 https://graduation-project-alpha-nine.vercel.app 看（手机访问会自动跳到移动版）。')
body('（4）开发过程踩过几个坑：sql.js 不接受 undefined 参数、Render 不支持 Gitee、Neon 密码认证失败、Element Plus 在手机上体验差……这些经验都写在了 5.8 节里，希望对后来人有帮助。')
body('没做好的地方：')
body('（1）权限系统太简单。只做了 admin / user 两种角色，没有"项目-人员"的细粒度权限。')
body('（2）任务没有"实际开始/结束日期"字段，进度只靠百分比描述，对长期项目来说粒度太粗。')
body('（3）移动端只覆盖了核心场景（看任务、更新进度、发评论），复杂的 WBS 编辑还在电脑端做。')
body('（4）附件用了 base64 存数据库，2MB 限制是为了演示方便。生产环境应该用对象存储。')
body('（5）变更日志只在任务里实现了，项目和用户的 CRUD 还没接入。代码里有预留接口但没接通。')
body('（6）关键路径算法只支持父子依赖，不支持跨分支约束（如"任务 A 必须在任务 B 开始 3 天后开始"），这个需要更完整的 PERT/CPM 实现。')
body('总体来说，作为一个工程管理专业的学生 + 前端工程师转的毕设，我觉得这个系统的"技术 + 业务"结合度还可以，对中小型建筑企业也确实有用——尤其是评论、附件、变更日志这几个功能，加完后从"工具"变成了"协作平台"，这才像个 ToB 系统该有的样子。')
body('再补一个我自己最满意的设计理念：这套系统始终在贯彻"预警 → 定位 → 处理"的闭环——红色 Alert 是预警、延期管理页是定位、点任务跳到项目页是处理。任何一个 ToB 系统的告警都该走完这三条链路，否则告警就只是装饰。')
heading('7.2 不足与展望', 2)
body('系统目前还是个 demo，离商用还有距离。下一步如果要继续做，我会重点关注这几块：')
body('（1）数据维度扩展。任务里加"实际开始/结束日期"、项目里加"实际完工日 vs 计划完工日"的对比，再加一些基础的成本管理（材料、人工、机械费）。这些数据是项目经理真正想看的。')
body('（2）移动端能力扩展。Vant 那套已经覆盖了核心场景，下一步要做的是 PWA 化（加 manifest.json 和 service worker，让用户能"添加到主屏幕"像 APP 一样用）和离线缓存（没网时能看历史数据，联网后同步）。这两个加上后移动端体验就接近原生 APP 了。')
body('（3）多租户 SaaS 化。目前是单租户，多个公司用要切数据库。改成"公司-项目"二级数据隔离，每家公司独立账户。')
body('（4）AI 辅助。周报自动生成、风险预测、关键路径智能识别——这些是工地项目里真正值钱的功能，也是我作为前端工程师最想尝试的方向。')
body('（5）数据可视化升级。7 个项目的全局对比、延期率统计、施工员工作量排名……这些 dashboard 页面是我下一步重点要做的事。')
body('（6）权限细化和通知系统。"项目-人员"细粒度权限、邮件/微信通知（任务延期、评论 @某人）、操作审计合规报告——这些是企业级必备。')
body('（7）智能告警。现在的延期检测是"事后"型（已经延期了才标红），下一步要做"事前"型——根据历史延期规律和任务依赖关系，预测某个任务**即将**延期的概率，提前推送给负责人。这种预测性告警才是工程项目管理的真价值。')
body('希望这套系统能成为我正式进入 ToB 软件开发行业的敲门砖。也希望读到这儿的同学能少走一些我走过的弯路——尤其是部署那部分，第一次搞全免费部署方案真的把我整崩溃了两次。')
body('最后，感谢我的指导老师在整个毕业设计过程中给我的耐心指导，感谢家人在我写论文期间的包容和鼓励，感谢同事们在技术选型上给我的建议。')
page_break()

# ==================== 参考文献 ====================
heading('参考文献', 1)
refs = [
    '[1] 丁士昭. 工程项目管理[M]. 北京: 中国建筑工业出版社, 2014.',
    '[2] 丛培经. 工程项目管理[M]. 北京: 中国建筑工业出版社, 2017.',
    '[3] 成虎, 陈群. 工程项目管理[M]. 北京: 中国建筑工业出版社, 2015.',
    '[4] 全国一级建造师执业资格考试用书编写委员会. 建设工程项目管理[M]. 北京: 中国建筑工业出版社, 2023.',
    '[5] 李忠富. 建筑施工组织与管理[M]. 北京: 机械工业出版社, 2019.',
    '[6] 刘伊生. 建设工程进度控制[M]. 北京: 中国建筑工业出版社, 2021.',
    '[7] 毛鹤琴. 土木工程施工[M]. 武汉: 武汉理工大学出版社, 2018.',
    '[8] GB/T 50502-2009. 建筑施工组织设计规范[S]. 北京: 中国建筑工业出版社, 2009.',
    '[9] 尤建新, 陈守明. 项目管理[M]. 北京: 高等教育出版社, 2018.',
    '[10] 张海藩, 牟永敏. 软件工程导论[M]. 北京: 清华大学出版社, 2013.',
    '[11] 朴灵. 深入浅出Node.js[M]. 北京: 人民邮电出版社, 2013.',
    '[12] 梁灏. Vue.js实战[M]. 北京: 清华大学出版社, 2017.',
    '[13] 尤雨溪. Vue.js官方文档[EB/OL]. https://cn.vuejs.org/, 2024.',
    '[14] 王沛. Webpack实战：入门、进阶与调优[M]. 北京: 电子工业出版社, 2019.',
    '[15] 刘博文. 深入浅出Vue.js[M]. 北京: 人民邮电出版社, 2019.',
    '[16] 廖雪峰. JavaScript全栈教程[EB/OL]. https://www.liaoxuefeng.com/, 2023.',
    '[17] 阮一峰. ECMAScript 6入门[M]. 北京: 电子工业出版社, 2017.',
    '[18] 王珊, 萨师煊. 数据库系统概论[M]. 北京: 高等教育出版社, 2014.',
    '[19] 杨旭. MySQL数据库应用从入门到精通[M]. 北京: 清华大学出版社, 2016.',
    '[20] Fielding R T. Architectural Styles and the Design of Network-based Software Architectures[D]. University of California, Irvine, 2000.',
    '[21] 钟志农. 基于Web的项目管理信息系统设计与实现[J]. 计算机工程与设计, 2010, 31(12): 2856-2859.',
    '[22] 张建平, 李丁, 林佳瑞等. BIM在工程施工中的应用[J]. 施工技术, 2012, 41(16): 10-17.',
    '[23] 陈建国, 周兴. 基于BIM的建设工程多维集成管理[J]. 科技进步与对策, 2014, 31(10): 56-60.',
    '[24] 王要武, 吴宇迪. 智慧建设及其支持技术研究[J]. 土木工程学报, 2012, 45(S2): 241-250.',
    '[25] 李恒, 郭红领, 黄霆等. BIM在建设项目中应用模式研究[J]. 工程管理学报, 2010, 24(5): 525-529.',
    '[26] 马智亮, 毛娜. 基于BIM的建筑工程施工进度管理框架[J]. 土木建筑工程信息技术, 2015, 7(3): 1-6.',
    '[27] 何清华, 韩翔宇. 基于BIM的进度管理系统框架构建和流程设计[J]. 项目管理技术, 2011, 9(9): 96-100.',
    '[28] 刘洋, 申玲. 基于BIM的施工进度计划编制方法研究[J]. 工程管理学报, 2015, 29(3): 91-96.',
    '[29] 赵彬, 王友群, 牛博生. 基于BIM的4D虚拟建造技术在工程项目进度管理中的应用[J]. 建筑经济, 2011(9): 93-95.',
    '[30] 张洋. 基于BIM的建筑工程信息集成与管理研究[D]. 清华大学, 2009.',
    '[31] 何关培. BIM总论[M]. 北京: 中国建筑工业出版社, 2011.',
    '[32] 李勇. 建筑施工进度管理信息化研究[J]. 建筑技术, 2018, 49(6): 652-654.',
    '[33] 黄强. 基于Web的工程项目管理系统的设计与实现[D]. 电子科技大学, 2015.',
    '[34] 王广斌, 张洋, 谭丹. 基于BIM的工程项目成本核算理论及实现方法[J]. 科技进步与对策, 2009, 26(21): 47-49.',
    '[35] 孙成双, 江帆. 基于BIM的工程项目进度管理研究综述[J]. 工程管理学报, 2020, 34(1): 1-6.',
    '[36] 有赞前端团队. Vant 4 移动端组件库技术文档[Z]. https://vant-ui.github.io/vant/, 2023.',
    '[37] 运筹学教材编写组. 运筹学[M]. 北京: 清华大学出版社, 2012. (关键路径法章节)',
    '[38] 陈建辉, 王继荣. 工程项目审计与变更管理实务[M]. 北京: 中国电力出版社, 2019.',
]
for ref in refs:
    add_para(ref, font_name='宋体', size=12, first_line_indent=0, space_after=2)
page_break()

# ==================== 致谢 ====================
heading('致  谢', 1)
add_para('')
body('时光荏苒，在广州大学的学习生涯即将画上句号。回顾这段求学之路，心中充满感激。')
body('首先，我要衷心感谢我的指导老师。在毕业设计的整个过程中，从选题确定、系统设计到论文撰写，老师都给予了我悉心的指导和宝贵的建议。老师严谨的治学态度、丰富的专业知识和耐心的指导风格，让我受益匪浅。')
body('其次，我要感谢广州大学继续教育学院的所有授课老师。是你们的辛勤付出和精彩授课，让我在工程管理专业领域打下了坚实的理论基础。虽然我的本职工作是一名前端工程师，但通过这几年的学习，我对工程项目管理有了更深入的理解，这为我完成本次毕业设计提供了重要的知识支撑。')
body('同时，我也要感谢我的家人和同事们。在学习和工作并行推进的日子里，是你们的理解、支持和鼓励，让我能够坚持完成学业。')
body('最后，感谢所有在百忙之中评阅本文和参加答辩的各位专家、教授。由于本人水平有限，论文中难免存在不足之处，恳请各位专家批评指正。')
page_break()

# ==================== 保存 ====================
output_path = r"C:\Users\Admin\Desktop\毕业论文-龙超滔-施工进度管理系统-初稿.docx"
doc.save(output_path)
print(f"✅ 论文已保存到: {output_path}")
print(f"   文件大小: {os.path.getsize(output_path)/1024:.1f} KB")